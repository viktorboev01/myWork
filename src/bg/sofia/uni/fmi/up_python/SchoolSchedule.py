import xlsxwriter
import docx
from docx.enum.table import WD_ROW_HEIGHT_RULE
from datetime import time
from enum import Enum
import itertools
import re
import os

class Days(Enum):
    Monday = 1
    Tuesday = 2
    Wensday = 3
    Thursday = 4
    Friday = 5
    Saturday = 6
    Sunday = 7

class Formats(Enum):
    group_format = {
                'bold': 1,
                'border': 1,
                'align': 'center',
                'valign': 'vcenter',
                'fg_color': '#ffb6c1'}
    interval_format = {
                'bold': 1,
                'border': 1,
                'align': 'center',
                'valign': 'vcenter',
                'fg_color': '#ADD8E6',
                'num_format': 'hh:mm'}
    subject_format = {
                'bold': 1,
                'border': 1,
                'align': 'center',
                'valign': 'vcenter',
                'fg_color': '#90EE90'}
    teacher_format = {
                'bold': 1,
                'border': 1,
                'align': 'center',
                'valign': 'vcenter',
                'fg_color': '#FAFA33'}
    day_format = {
                'bold': 1,
                'border': 1,
                'align': 'center',
                'valign': 'vcenter',
                'fg_color': '#ADD8E6'}

class IntervalException(Exception):
    pass

class GroupException(Exception):
    pass

class SubjectException(Exception):
    pass

class TeacherException(Exception):
    pass

class InstitutionDataException(Exception):
    pass

class ProgramException(Exception):
    pass

class Interval:

    def __init__(self, day, start, num_hour):
        
        self.init_validation(day, start, num_hour)

        self.day = day
        self.start = start
        self.num_hour = num_hour

        self.teachers = set()
        self.groups = {}

    def init_validation(self, day, start, num_hour):
        
        if type(start) != time:
            raise IntervalException("Parameter 'start' has incorrect value")
        if type(day) != Days:
            raise IntervalException("Parameter 'day' has incorrect value")
        if type(num_hour) != int or num_hour < 0:
            raise IntervalException("Parameter 'num_hour' has incorrect value")

    def __str__(self):
        return str(self.day) + ' ' + str(self.start)

    def __hash__(self):
        return hash(str(self))

    def _compare(self, other):
        if self.day.value < other.day.value:
            return -1
        if self.day.value > other.day.value:
            return 1
        if self.start < other.start:
            return -1
        if self.start > other.start:
            return 1
        return 0

    def __lt__(self, other):
        return self._compare(other) < 0

    def __eq__(self, other):
        return self._compare(other) == 0

    def __ne__(self, other):
        return self._compare(other) != 0

    def __gt__(self, other):
        return self._compare(other) > 0

class Teacher:

    id_iter = itertools.count()

    def __init__(self, first_name, last_name, subject):

        self.init_validation(first_name, last_name, subject)

        self.id = next(self.id_iter)
        self.first_name = first_name
        self.last_name = last_name
        self.personal_intervals = set()
        self.subject = subject
        
        subject.add_teacher(self)

    def init_validation(self, first_name, last_name, subject):
        
        if type(first_name) != str or not re.search("^[A-Z][a-z]*$", first_name):
            raise TeacherException("Parameter 'first_name' has incorrect value")
        if type(last_name) != str or not re.search("^[A-Z][a-z]*$", last_name):
            raise TeacherException("Parameter 'last_name' has incorrect value")
        if type(subject) != Subject:
            raise TeacherException("Parameter 'subject' has incorrect value")

    def __eq__(self, other):
        return self.id == other.id

    def __ne__(self, other):
        return not self.__eq__(other)

    def add_personal_interval(self, interval):
        if type(interval) != Interval:
            raise IntervalException()
        self.personal_intervals.add(interval)

    def __str__(self):
        return self.first_name + ' ' + self.last_name 

    def __hash__(self):
        return hash(str(self))

class Subject:

    def __init__(self, name, level):

        self.init_validation(name, level)

        self.name = name
        self.level = level
        self.teachers = set()

    def init_validation(self, name, level):
        
        if type(name) != str or not re.search("^[A-Z]?[a-z]*(\s[0-9])?$", name):
            raise SubjectException("Parameter 'name' has incorrect value")
        if type(level) != int:
            raise SubjectException("Parameter 'level' has incorrect value")

    def add_teacher(self, teacher):
        self.teachers.add(teacher)

    def __str__(self):
        return self.name 

    def __hash__(self):
        return hash(str(self))

    def __eq__(self, other):
        return self.name.__eq__(other.name) and self.level == other.level

    def __ne__(self, other):
        return not self.__eq__(other)

class InstitutionData:

    def __init__(self, teachers, subjects, days, start_time, end_time, number_of_levels, duration_lesson):

        self.teachers = set()
        self.subjects = set()
        self.days = []

        self.init_validation(teachers, subjects, days, start_time, end_time, number_of_levels, duration_lesson)
        self.number_of_levels = number_of_levels

        self.all_intervals = []
        self._get_all_intervals(start_time, end_time, duration_lesson)
        self.all_times_per_day = self._get_times_per_day(start_time, end_time, duration_lesson)

    def _add_minutes_type_time(self, hours, minutes, minutes_to_add):
        while minutes + minutes_to_add >= 60:
            hours += 1
            minutes_to_add -= 60
        minutes = minutes + minutes_to_add
        return time(hours, minutes)

    def _get_times_per_day(self, start_time, end_time, duration_lesson):
        intervals_per_day = []
        current_time = start_time
        while current_time != end_time:
            intervals_per_day.append(current_time)
            current_time = self._add_minutes_type_time(current_time.hour, current_time.minute, duration_lesson)
        return intervals_per_day

    def _get_all_intervals(self, start_time, end_time, duration_lesson):
        intervals_per_day = self._get_times_per_day(start_time, end_time, duration_lesson)
        for day in self.days:
            counter = 0
            for interval in intervals_per_day:
                self.all_intervals.append(Interval(day, interval, counter))
                counter += 1

    def init_validation(self, teachers, subjects, days, start_time, end_time, number_of_levels, minutes_per_lesson):
        
        if not isinstance(teachers, set):
            raise InstitutionDataException("Parameter 'teachers' has incorrect value")

        if not isinstance(subjects, set):
            raise InstitutionDataException("Parameter 'subjects' has incorrect value")

        if not isinstance(days, list):
            raise InstitutionDataException("Parameter 'days' has incorrect value")

        if type(start_time) != time:
            raise InstitutionDataException("Parameter 'start_time' has incorrect value")

        if type(end_time) != time:
            raise InstitutionDataException("Parameter 'end_time' has incorrect value")

        if type(number_of_levels) != int or number_of_levels <= 0:
            raise InstitutionDataException("Parameter 'number_of_levels' has incorrect value")

        if type(minutes_per_lesson) != int or minutes_per_lesson <= 0:
            raise InstitutionDataException("Parameter 'duration_lesson' has incorrect value")

        for teacher in teachers:
            if type(teacher) != Teacher:
                raise TeacherException("Element in 'teachers' from 'InstitutionData' has incorrect value")
            self.teachers.add(teacher)

        for subject in subjects:
            if type(subject) != Subject or subject.level > number_of_levels:
                raise SubjectException("Element in 'subjects' from 'InstitutionData' has incorrect value")
            self.subjects.add(subject)

        for day in days:
            if type(day) != Days:
                raise Exception("Element in 'days' from 'InstitutionData' has incorrect value")
            self.days.append(day)

class Group:

    def __init__(self, level, name, subjects):

        self.subjects = {}
        self.init_validation(level, name, subjects)
            
        self.level = level
        self.name = name

    def _compare(self, other):
        if (self.level != other.level):
            return self.level - other.level
        if self.name < other.name:
            return -1
        if self.name > other.name:
            return 1
        return 0

    def __lt__(self, other):
        return self._compare(other) < 0

    def __eq__(self, other):
        return self._compare(other) == 0

    def __ne__(self, other):
        return self._compare(other) != 0

    def __gt__(self, other):
        return self._compare(other) > 0

    def __str__(self):
        return str(self.level) + ' ' + self.name

    def __hash__(self):
        return hash(str(self))

    def init_validation(self, level, name, subjects):
        
        if not isinstance(subjects, dict):
            raise GroupException("Parameter 'subjects' has incorrect value")

        if type(level) != int or level <= 0:
            raise GroupException("Parameter 'level' has incorrect value")

        if type(name) != str:
            raise GroupException("Parameter 'name' has incorrect value")

        for subject in subjects:

            if type(subject) != Subject:
                raise SubjectException("Element in 'subjects' from 'Group' has incorrect value")

            if subject.level > level:
                raise GroupException("Element in 'subjects' from 'Group' has incorrect level for the group")

            times_a_week = subjects[subject]

            if type(times_a_week) != int or times_a_week <= 0:
                raise GroupException("Element in 'subjects' from 'Group' has incorrect times a week")

            self.subjects[subject] = times_a_week

class Program:

    def __init__(self, groups, institution):

        self.init_validation(groups, institution)

        self.groups = list(groups)
        self.groups.sort()
        self.teachers = institution.teachers
        self.subjects = institution.subjects
        self.all_intervals = institution.all_intervals
        self.all_times_per_day = institution.all_times_per_day
        self.days = institution.days

    def init_validation(self, groups, institution):

        if not isinstance(groups, set):
            raise ProgramException("Parameter 'groups' has incorrect value")

        if type(institution) != InstitutionData:
            raise ProgramException("Parameter 'institution' has incorrect value")

        for group in groups:
            if type(group) != Group:
                raise GroupException("Element in 'groups' from 'Program' has incorrect value")

            if group.level > institution.number_of_levels:
                raise ProgramException("Element in 'groups' from 'Program' has incorrect level for the institution")

    def add_to_schedule(self, interval, teacher, group):
        if teacher not in interval.teachers and group not in interval.groups:
            interval.teachers.add(teacher)
            interval.groups[group] = teacher
            return True
        return False

    def change_subject(self, interval, group, teacher):

        interval = self._get_interval(interval)
        if group in interval.groups and teacher not in interval.teachers:

            interval.teachers.remove(interval.groups[group])
            interval.groups[group] = teacher
            interval.teachers.add(teacher)
            return True

        return self.add_to_schedule(interval, teacher, group)

    def remove_from_schedule(self, interval, teacher, group):
        interval.teachers.remove(teacher)
        interval.groups.pop(group)

    def _search_interval_to_add(self, subject, group):
        for teacher in subject.teachers:
            for interval in self.all_intervals: 
                if self.add_to_schedule(interval, teacher, group):
                    return
        raise ProgramException('Not possible to make a program')

    def _prepare_list(self, list_subjects, dict_subjects):
        list_subjects.clear()
        stop = False
        while not stop:
            stop = True
            for subject, key in dict_subjects.items():
                if key != 0:
                    list_subjects.append(subject)
                    stop = False
                    dict_subjects[subject] = key - 1 

    def make_program(self):
        
        list_subjects = []
        temp_subjects = []
        current_counter = 0
        main_counter = 0
        for group in self.groups:
            self._prepare_list(list_subjects, dict_subjects = group.subjects.copy())
            for subject in list_subjects:
                if current_counter >= main_counter:
                    self._search_interval_to_add(subject, group)
                else:
                    temp_subjects.append(subject)
                current_counter += 1

            for subject in temp_subjects:
                self._search_interval_to_add(subject, group)

            temp_subjects.clear()
            current_counter = 0
            main_counter += 1

    def _get_interval(self, interval):
        for current in self.all_intervals:
            if current == interval:
                return current
        raise IntervalException()

    def move_unit(self, interval_from, interval_to, group):
        interval_from = self._get_interval(interval_from)
        interval_to = self._get_interval(interval_to)
        if group not in interval_from.groups:
            return False
        teacher = interval_from.groups[group]
        if teacher not in interval_to.teachers and group not in interval_to.groups:
            self.add_to_schedule(interval_to, teacher, group)
            self.remove_from_schedule(interval_from, teacher, group)
            return True
        return False

    def switch_two_units(self, interval1, group1, interval2, group2):

        interval1 = self._get_interval(interval1)
        interval2 = self._get_interval(interval2)

        if group1 not in interval1.groups and group2 not in interval2.groups:
            return False

        elif group1 not in interval1.groups:
            self.move_unit(interval2, interval1, group2)

        elif group2 not in interval2.groups:
            self.move_unit(interval1, interval2, group1)
        
        else:   
            teacher1 = interval1.groups[group1]
            teacher2 = interval2.groups[group2]

            if teacher2 not in interval1.teachers and 
            group2 not in interval1.groups and 
            teacher1 not in interval2.teachers and 
            group1 not in interval2.groups:

                self.move_unit(interval1, interval2, group1)
                self.move_unit(interval2, interval1, group2)
                return True

        return False

    def balance(self, group):
        intervals_to_move = []
        averageLength = (self.count_units_per_group(group) / len(self.days)) 
        for interval in self.all_intervals:
            
            if interval.num_hour >= averageLength and \
            group in interval.groups:
                intervals_to_move.append(interval)

            elif interval.num_hour < averageLength and \
            group not in interval.groups and \
            len(intervals_to_move) != 0:
                for i in intervals_to_move:
                    if self.move_unit(i, interval, group):
                        intervals_to_move.remove(i)
                        break
        return

    def count_units_per_group(self, group):
        counter = 0
        for interval in self.all_intervals:
            if group in interval.groups:
                counter += 1
        return counter
    
    def count_units(self):
        counter = 0
        for group in self.groups:
            counter += self.count_units_per_group(group)
        return counter

class ExcelExporter(Program):

    def __init__(self, groups, institution, path):

        Program.__init__(self, groups, institution)
        self.workbook = xlsxwriter.Workbook(path)
        self.worksheet = self.workbook.add_worksheet()
        self.path = path

        self.group_format = self.workbook.add_format(Formats.group_format.value)
        self.interval_format = self.workbook.add_format(Formats.interval_format.value)
        self.subject_format = self.workbook.add_format(Formats.subject_format.value)
        self.teacher_format = self.workbook.add_format(Formats.teacher_format.value)
        self.day_format = self.workbook.add_format(Formats.day_format.value)
    
    def _set_header(self, current_row, group):
        self.worksheet.write('A' + str(current_row), str(group), self.group_format)
        current_letter = 'B'
        current_row += 1
        str_current_row = str(current_row)
        for day in self.days:
            merge_string = current_letter + str_current_row + ':' + chr(ord(current_letter) + 1) + str_current_row
            self.worksheet.merge_range(merge_string, day.name, self.day_format)
            current_letter = chr(ord(current_letter) + 2)

    def _set_table(self):
        current_row = 1
        for group in self.groups:
            self._set_header(current_row, group)
            self._make_column_from_objects('A', current_row + 2, self.all_times_per_day)
            current_row += len(self.all_times_per_day) + 3

    def export(self):

        self.worksheet.set_column('A:Z', 12)
        self._set_table()
        current_letter = 'B'
        current_row = 3

        for group in self.groups:
            for interval in self.all_intervals:
                if group in interval.groups:
                    
                    letter = chr(ord(current_letter) + (interval.day.value - 1) * 2)
                    row = interval.num_hour

                    teacher = interval.groups.get(group)
                    subject = teacher.subject
                    self._make_column_from_objects(letter, current_row + row, [subject]) 
                    self._make_column_from_objects(chr(ord(letter) + 1), current_row + row, [teacher])

            current_row += len(self.all_times_per_day) + 3
            
        self.workbook.close()
        os.system("start " + self.path)

    def _prepare_data(self, object):
            
        if type(object) == time:
            return object.strftime('%H:%M'), self.interval_format

        if type(object) == Subject:
            return object.name, self.subject_format
        
        else:
            return object.last_name, self.teacher_format

    def _make_column_from_objects(self, start_letter, start_row, objects):
        counter = start_row
        for object in objects:
            data_and_format = self._prepare_data(object)
            self.worksheet.write(start_letter + str(counter), data_and_format[0], data_and_format[1])
            counter += 1

class DocsExporter(Program):
    
    def __init__(self, groups, institution, path):
        Program.__init__(self, groups, institution) 
        self.path = path 
        self.rows = []

    def _make_header(self, group_table):
        
        cells = group_table.rows[0].cells
        counter = 1
        for day in self.days:
            cells[counter].text = day.name
            counter += 1

    def _set_table(self, group_table):
        
        self._make_header(group_table)
        counter = 1

        for interval in self.all_times_per_day:
            cell = group_table.rows[counter].cells[0]
            cell.text = interval.strftime('%H:%M')
            counter += 1

    def export(self):

        doc = docx.Document()
        doc.add_heading('School Schedule',0)
        
        for group in self.groups:

            doc.add_paragraph('\n' + str(group))
            group_table = doc.add_table(rows = len(self.all_times_per_day) + 1, cols = len(self.days) + 1)
            group_table.style= "Table Grid"
            for row in group_table.rows:
                row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
            self._set_table(group_table)

            for interval in self.all_intervals:
                if group in interval.groups:
                    column = interval.day.value
                    subject = interval.groups.get(group).subject.name
                    teacher = interval.groups.get(group).last_name
                    data = subject + ', ' + teacher
                    group_table.rows[interval.num_hour + 1].cells[column].text = data

        doc.save(self.path)
        os.system("start " + self.path)
        
def test():

    days = [Days.Monday, Days.Tuesday, Days.Wensday, Days.Thursday, Days.Friday, Days.Saturday]
            
    level = 3
    subject1 = Subject('geography', level)
    subject2 = Subject('maths', level)
    subject3 = Subject('sport', level)

    dict_subjects = {subject1:4, subject2:2, subject3:1}
    dict_subjects1 = {subject1:2, subject2:2, subject3:5}
    dict_subjects2 = {subject1:2, subject2:2, subject3:2}
    set_subjects = {subject1, subject2, subject3}

    teacher1 = Teacher('Ivan', 'Ivanov', subject1)
    teacher2 = Teacher('Petar', 'Petrov', subject2)
    teacher3 = Teacher('Georgi', 'Georgiev', subject3)
    teacher4 = Teacher('Ivan', 'Hristov', subject1)
    teacher5 = Teacher('Petar', 'Trendafilov', subject2)
    teacher6 = Teacher('Georgi', 'Gochev', subject3)

    group1 = Group(level, 'a', dict_subjects)
    group2 = Group(level, 'b', dict_subjects1)
    group3 = Group(level, 'c', dict_subjects2)

    set_teachers = {teacher1, teacher2, teacher3, teacher4, teacher5, teacher6}
    set_subjects = {subject1, subject2, subject3}
    set_groups = {group1, group2, group3}

    institution = InstitutionData(set_teachers, set_subjects, days, time(8, 00), time(11, 00), level, 60)

    #exporter = DocsExporter(set_groups, institution, 'C:/homework_for_school/Python/Program.doc')
    exporter = ExcelExporter(set_groups, institution, 'C:/homework_for_school/Python/Program.xlsx')
    i1 = Interval(Days.Wensday, time(8,00), 0)
    i2 = Interval(Days.Thursday, time(8,00), 0)
    exporter.make_program()
    #print(exporter.move_unit(i1, i2, group2))
    #print(exporter.switch_two_units(i2, group2, i1, group1))
    #print(exporter.change_subject(i2, group1, teacher2))
    #print(exporter.change_subject(i1, group2, teacher2))
    #exporter.balance(group1)
    #exporter.balance(group2)
    #exporter.balance(group3)
    exporter.export()

test()